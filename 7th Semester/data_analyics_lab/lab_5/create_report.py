import nbformat
from docx import Document

def create_docx_from_notebook(notebook_path, docx_path):
    # Load the Jupyter notebook
    with open(notebook_path, 'r', encoding='utf-8') as nb_file:
        notebook = nbformat.read(nb_file, as_version=4)

    # Create a new Document for the Word file
    doc = Document()
    doc.add_heading('Jupyter Notebook Code and Outputs', 0)

    # Iterate through each cell in the notebook
    for cell in notebook.cells:
        if cell.cell_type == 'code':
            # Add code to document
            doc.add_heading('Code:', level=1)
            doc.add_paragraph(cell.source)

            # Add output to document (if present)
            if 'outputs' in cell:
                for output in cell.outputs:
                    if hasattr(output, 'data') and 'text/plain' in output.data:
                        # Add output if it's of type text/plain
                        doc.add_heading('Output:', level=2)
                        doc.add_paragraph(output.data['text/plain'])
                    elif hasattr(output, 'text'):
                        # Handle legacy output with `text` attribute
                        doc.add_heading('Output:', level=2)
                        doc.add_paragraph(output.text)

        elif cell.cell_type == 'markdown':
            # Add markdown (text) to document
            doc.add_heading('Text:', level=1)
            doc.add_paragraph(cell.source)

    # Save the document as .docx
    doc.save(docx_path)
    print(f'Document saved as {docx_path}')

# Example usage
notebook_path = 'task.ipynb'  # Replace with your notebook file path
docx_path = 'da_lab_05.docx'  # Replace with the desired output docx file path
create_docx_from_notebook(notebook_path, docx_path)
